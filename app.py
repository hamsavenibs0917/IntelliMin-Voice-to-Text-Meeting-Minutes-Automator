"""
🎙️ IntelliMin : Voice-to-Text Meeting Minutes Automator
Standalone Streamlit Application with Clean Text Preprocessing, TF-IDF Naive Bayes Model, WordCloud & Dual Charts
"""

import os
import re
import json
import sqlite3
import tempfile
from datetime import datetime
from io import BytesIO

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
from groq import Groq

# Imports Scikit-Learn ML modules
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.naive_bayes import MultinomialNB

# Imports ReportLab for PDF
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable

# Imports python-docx for DOCX
import docx

DB_PATH = "intelliminutes.db"
MAX_STORED_MEETINGS = 10

# Global trained ML objects
_vectorizer = None
_classifier_model = None


def clean_text(text: str) -> str:
    # Preprocesses and cleans text
    text = str(text).lower()
    text = re.sub(r"[^a-zA-Z\s]", "", text)
    return text


def train_nlp_intent_model():
    # Trains TF-IDF Naive Bayes
    global _vectorizer, _classifier_model
    if os.path.exists("meeting_dataset.csv"):
        df_train = pd.read_csv("meeting_dataset.csv")
    else:
        data = {
            "Text": [
                "We decided to approve the project budget.",
                "Rahul will complete the report by Friday.",
                "We discussed the new project requirements.",
                "Priya will prepare the presentation.",
                "The team decided to use Python for the project.",
                "We discussed the problems in the current system.",
                "David will submit the final report tomorrow.",
                "The team discussed the project deadline."
            ],
            "Category": [
                "Decision",
                "Action Item",
                "Discussion",
                "Action Item",
                "Decision",
                "Discussion",
                "Action Item",
                "Discussion"
            ]
        }
        df_train = pd.DataFrame(data)

    df_train["Clean_Text"] = df_train["Text"].apply(clean_text)
    _vectorizer = TfidfVectorizer()
    X_train = _vectorizer.fit_transform(df_train["Clean_Text"])
    y_train = df_train["Category"]

    _classifier_model = MultinomialNB()
    _classifier_model.fit(X_train, y_train)


# Trains model on startup
train_nlp_intent_model()


def predict_sentence_intent(text: str) -> str:
    # Predicts sentence intent category
    global _vectorizer, _classifier_model
    if _vectorizer is None or _classifier_model is None:
        train_nlp_intent_model()
    try:
        clean_t = clean_text(text)
        vec = _vectorizer.transform([clean_t])
        pred = _classifier_model.predict(vec)[0]
        return str(pred)
    except Exception:
        return "Discussion"


def get_groq_api_key() -> str:
    # Retrieves Groq API key from OS environment, .env file, or Streamlit Secrets
    key = os.environ.get("GROQ_API_KEY", "").strip()
    if not key and hasattr(st, "secrets"):
        try:
            if "GROQ_API_KEY" in st.secrets:
                key = str(st.secrets["GROQ_API_KEY"]).strip()
        except Exception:
            pass
    return key


def get_groq_client():
    # Returns active Groq client using environment variable or Streamlit Secrets
    key = get_groq_api_key()
    if not key:
        st.warning("⚠️ **Groq API Key Missing**: Please configure `GROQ_API_KEY` in your `.env` file locally or in Streamlit Cloud Secrets.")
    return Groq(api_key=key if key else "missing_api_key")


def init_database():
    # Connects to SQLite database
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS meetings (
            meeting_id TEXT PRIMARY KEY,
            title TEXT,
            date TEXT,
            target_language TEXT,
            main_summary TEXT,
            off_topic_report TEXT,
            full_json_details TEXT,
            productivity_score REAL
        )
    """)

    cursor.execute("PRAGMA table_info(meetings)")
    columns = [info[1] for info in cursor.fetchall()]
    if "full_json_details" not in columns:
        cursor.execute("ALTER TABLE meetings ADD COLUMN full_json_details TEXT")

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS transcripts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            meeting_id TEXT,
            speaker TEXT,
            original_text TEXT,
            translated_text TEXT,
            language TEXT,
            category TEXT,
            is_off_topic INTEGER,
            topic TEXT,
            FOREIGN KEY (meeting_id) REFERENCES meetings (meeting_id)
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS action_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            meeting_id TEXT,
            task TEXT,
            assignee TEXT,
            deadline TEXT,
            status TEXT,
            FOREIGN KEY (meeting_id) REFERENCES meetings (meeting_id)
        )
    """)

    conn.commit()
    conn.close()
    enforce_10_meeting_cap()


def enforce_10_meeting_cap():
    # Enforces ten meeting limit
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM meetings WHERE meeting_id NOT IN (SELECT meeting_id FROM meetings ORDER BY date DESC LIMIT ?)", (MAX_STORED_MEETINGS,))
    cursor.execute("DELETE FROM transcripts WHERE meeting_id NOT IN (SELECT meeting_id FROM meetings)")
    cursor.execute("DELETE FROM action_items WHERE meeting_id NOT IN (SELECT meeting_id FROM meetings)")
    conn.commit()
    conn.close()


def save_meeting_to_db(meeting_id: str, title: str, date_str: str, target_lang: str, utterances: list, summary: str, off_topic_rep: str, action_items: list, full_details: dict = None):
    # Saves meeting into SQLite
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    json_str = json.dumps(full_details, ensure_ascii=False) if full_details else "{}"
    cursor.execute("""
        INSERT OR REPLACE INTO meetings (meeting_id, title, date, target_language, main_summary, off_topic_report, full_json_details, productivity_score)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, (meeting_id, title, date_str, target_lang, summary, off_topic_rep, json_str, 90.0))

    cursor.execute("DELETE FROM transcripts WHERE meeting_id = ?", (meeting_id,))
    cursor.execute("DELETE FROM action_items WHERE meeting_id = ?", (meeting_id,))

    for u in utterances:
        spk = u.get("speaker", "Person 1 (Detected Voice)")
        if not spk or str(spk).strip().lower() in ["unknown", "none", "nan", ""]:
            spk = "Person 1 (Detected Voice)"

        cursor.execute("""
            INSERT INTO transcripts (meeting_id, speaker, original_text, translated_text, language, category, is_off_topic, topic)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            meeting_id,
            spk,
            u.get("original_text", ""),
            u.get("translated_text", u.get("original_text", "")),
            u.get("language", u.get("detected_language", "English")),
            u.get("category", "Discussion"),
            1 if u.get("is_off_topic") else 0,
            u.get("topic", "General")
        ))

    for a in action_items:
        assignee = a.get("assignee", "Person 1 (Detected Voice)")
        if not assignee or str(assignee).strip().lower() in ["unknown", "none", "nan", ""]:
            assignee = "Person 1 (Detected Voice)"

        cursor.execute("""
            INSERT INTO action_items (meeting_id, task, assignee, deadline, status)
            VALUES (?, ?, ?, ?, ?)
        """, (
            meeting_id,
            a.get("task", ""),
            assignee,
            a.get("deadline", "Not Specified"),
            a.get("status", "Pending Follow-up")
        ))

    conn.commit()
    conn.close()
    enforce_10_meeting_cap()


def delete_meeting_db(meeting_id: str) -> bool:
    # Deletes single meeting record
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("DELETE FROM meetings WHERE meeting_id = ?", (meeting_id,))
        cursor.execute("DELETE FROM transcripts WHERE meeting_id = ?", (meeting_id,))
        cursor.execute("DELETE FROM action_items WHERE meeting_id = ?", (meeting_id,))
        conn.commit()
        conn.close()
        return True
    except Exception:
        return False


def search_meetings_db(query: str = "") -> list:
    # Searches stored meeting records
    enforce_10_meeting_cap()
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    if not query or len(query.strip()) == 0:
        cursor.execute("SELECT meeting_id, title, date, target_language, main_summary, off_topic_report FROM meetings ORDER BY date DESC LIMIT 10")
    else:
        q = f"%{query.strip()}%"
        cursor.execute("""
            SELECT meeting_id, title, date, target_language, main_summary, off_topic_report 
            FROM meetings 
            WHERE title LIKE ? OR date LIKE ? OR main_summary LIKE ? OR off_topic_report LIKE ?
            ORDER BY date DESC LIMIT 10
        """, (q, q, q, q))

    rows = cursor.fetchall()
    conn.close()

    results = []
    for r in rows:
        results.append({
            "meeting_id": r[0],
            "title": r[1],
            "date": r[2],
            "target_language": r[3],
            "main_summary": r[4],
            "off_topic_report": r[5]
        })
    return results


def get_meeting_details_db(meeting_id: str) -> dict:
    # Fetches full meeting details
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute("SELECT meeting_id, title, date, target_language, main_summary, off_topic_report, full_json_details FROM meetings WHERE meeting_id = ?", (meeting_id,))
    m_row = cursor.fetchone()
    if not m_row:
        conn.close()
        return {}

    full_details = {}
    if len(m_row) > 6 and m_row[6]:
        try:
            full_details = json.loads(m_row[6])
        except Exception:
            full_details = {}

    cursor.execute("SELECT speaker, original_text, translated_text, language, category, is_off_topic, topic FROM transcripts WHERE meeting_id = ?", (meeting_id,))
    t_rows = cursor.fetchall()
    transcripts = []
    for t in t_rows:
        transcripts.append({
            "speaker": t[0],
            "original_text": t[1],
            "translated_text": t[2],
            "language": t[3],
            "category": t[4],
            "is_off_topic": bool(t[5]),
            "topic": t[6]
        })

    cursor.execute("SELECT task, assignee, deadline, status FROM action_items WHERE meeting_id = ?", (meeting_id,))
    a_rows = cursor.fetchall()
    action_items = []
    for a in a_rows:
        action_items.append({
            "task": a[0],
            "assignee": a[1],
            "deadline": a[2],
            "status": a[3]
        })

    conn.close()

    res = {
        "meeting_id": m_row[0],
        "title": m_row[1],
        "date": m_row[2],
        "target_language": m_row[3],
        "main_summary": m_row[4],
        "off_topic_report": m_row[5],
        "transcripts": transcripts,
        "action_items": action_items
    }
    if full_details:
        res["type_a_report"] = full_details.get("type_a_report", {})
    return res


def transcribe_audio_groq(audio_bytes: bytes, filename: str = "input_audio.wav") -> dict:
    # Transcribes audio with Groq Whisper ASR
    if not audio_bytes or len(audio_bytes) == 0:
        return {"error": "Empty audio file received. Please select or record a valid audio file.", "text": ""}

    client = get_groq_client()
    ext = os.path.splitext(filename)[1] or ".wav"
    if not ext.startswith("."):
        ext = "." + ext

    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp_file:
        tmp_file.write(audio_bytes)
        tmp_path = tmp_file.name

    whisper_models = ["whisper-large-v3", "whisper-large-v3-turbo"]
    last_err = None

    try:
        for w_model in whisper_models:
            try:
                with open(tmp_path, "rb") as audio_file:
                    transcription = client.audio.transcriptions.create(
                        file=(filename, audio_file),
                        model=w_model,
                        response_format="verbose_json"
                    )

                text = getattr(transcription, "text", "")
                language = getattr(transcription, "language", "Unknown")
                
                if text and len(text.strip()) > 0:
                    return {
                        "text": text.strip(),
                        "language": language.capitalize() if isinstance(language, str) else "Detected Speech"
                    }
            except Exception as ex:
                last_err = ex
                continue

        if last_err:
            return {"error": f"Groq Whisper ASR error: {str(last_err)}", "text": ""}
        return {"error": "No speech recognized from audio clip.", "text": ""}
    finally:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass


def process_meeting_transcript_groq(text_or_utterances: list, target_language: str = "English") -> dict:
    # Evaluates Type A report using Groq LLM with multi-model fallback
    client = get_groq_client()
    input_data_str = json.dumps(text_or_utterances, ensure_ascii=False)

    prompt = f"""You are an expert Executive Assistant and Professional Meeting Minute Analyst.
Analyze the following meeting transcript / utterances:
{input_data_str}

CRITICAL SPEAKER NAMING RULES:
- If a person's actual name was explicitly stated in the discussion (e.g. "David here"), use their stated name.
- IF A PERSON'S NAME IS UNAVAILABLE OR UNMENTIONED: Mention them strictly as "Person 1 (Detected Voice)", "Person 2 (Detected Voice)", "Person 3 (Detected Voice)".
- DO NOT invent or guess arbitrary names (such as Rahul, Alice, John) if they were not stated in the input!

LANGUAGE REQUIREMENT:
- Detect input language and translate all summaries, arguments, outcomes, topics, and action items into {target_language}.

Output ONLY a valid JSON object matching this exact Type A Professional Meeting Report structure:
{{
  "discussion_title": "Specific Discussion Topic Title in {target_language}",
  "detected_language": "Detected Spoken Language",
  "type_a_report": {{
    "metadata": {{
      "meeting_title": "Title in {target_language}",
      "date": "{datetime.now().strftime('%Y-%m-%d')}",
      "time": "{datetime.now().strftime('%H:%M')}",
      "location": "Virtual Conference / Meeting Room",
      "facilitator": "Person 1 (Detected Voice)",
      "note_taker": "IntelliMin AI Automator",
      "attendees": ["Person 1 (Detected Voice)", "Person 2 (Detected Voice)"],
      "absentees": ["None Noted"]
    }},
    "executive_summary": {{
      "objective": "Core purpose or problem addressed in {target_language}",
      "outcome": "Brief two-to-three sentence summary of final results in {target_language}"
    }},
    "discussion_points": [
      {{
        "agenda_item": "Agenda Topic Name in {target_language}",
        "key_arguments": "Main perspectives, ideas, and data presented in {target_language}",
        "consensus": "Points where participants agreed or disagreed in {target_language}"
      }}
    ],
    "action_items": [
      {{
        "task": "Specific deliverable in {target_language}",
        "assignee": "Person 1 (Detected Voice) / Stated Name",
        "deadline": "Clear due date or Next Friday",
        "status": "Pending Follow-up"
      }}
    ],
    "future_planning": {{
      "open_issues": "Items tabled or postponed for later in {target_language}",
      "next_meeting": "Proposed date, time, and preliminary agenda in {target_language}"
    }}
  }},
  "processed_sentences": [
    {{
      "speaker": "Person 1 (Detected Voice) / Stated Name",
      "original_text": "...",
      "translated_text": "Sentence translated in {target_language}",
      "detected_language": "...",
      "is_off_topic": false,
      "topic": "Topic Name",
      "category": "Decision / Action Item / Discussion / Casual"
    }}
  ],
  "main_concept_summary": "Executive summary in {target_language}",
  "off_topic_report": "Summary of casual side conversations in {target_language}",
  "action_items": [
    {{
      "task": "Task description",
      "assignee": "Person 1 (Detected Voice)",
      "deadline": "Deadline",
      "status": "Pending"
    }}
  ]
}}
"""

    candidate_models = ["openai/gpt-oss-120b", "qwen/qwen3.6-27b", "openai/gpt-oss-20b", "groq/compound", "llama-3.3-70b-versatile"]
    last_error = None

    for model_name in candidate_models:
        try:
            response = client.chat.completions.create(
                model=model_name,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                response_format={"type": "json_object"}
            )
            data = json.loads(response.choices[0].message.content)
            
            for s in data.get("processed_sentences", []):
                txt = s.get("original_text", "")
                if txt:
                    s["category"] = predict_sentence_intent(txt)
                    if s["category"] == "Casual Off-Topic":
                        s["is_off_topic"] = True
                        
            return data
        except Exception as e:
            last_error = e
            continue

    # Fallback parsing using local NLP intent classifier if all LLM models fail
    processed = []
    for idx, u in enumerate(text_or_utterances if isinstance(text_or_utterances, list) else [text_or_utterances]):
        p_num = (idx % 3) + 1
        spk = f"Person {p_num} (Detected Voice)"
        utt_str = str(u)
        if ":" in utt_str:
            parts = utt_str.split(":", 1)
            spk = parts[0].strip()
            utt_str = parts[1].strip()
            
        cat_pred = predict_sentence_intent(utt_str)
        processed.append({
            "speaker": spk,
            "original_text": utt_str,
            "translated_text": utt_str,
            "detected_language": "Auto",
            "is_off_topic": True if cat_pred == "Casual Off-Topic" else False,
            "topic": "General Agenda",
            "category": cat_pred
        })
    return {
        "discussion_title": "Recorded Discussion Session",
        "detected_language": "Auto",
        "type_a_report": {
            "metadata": {
                "meeting_title": "Recorded Discussion Session",
                "date": datetime.now().strftime("%Y-%m-%d"),
                "time": datetime.now().strftime("%H:%M"),
                "location": "Virtual Conference",
                "facilitator": "Person 1 (Detected Voice)",
                "note_taker": "IntelliMin Automator",
                "attendees": ["Person 1 (Detected Voice)", "Person 2 (Detected Voice)"],
                "absentees": ["None"]
            },
            "executive_summary": {
                "objective": "Review key operational topics.",
                "outcome": "Session recorded and processed successfully."
            },
            "discussion_points": [
                {
                    "agenda_item": "General Discussion",
                    "key_arguments": "Captured spoken sentences.",
                    "consensus": "Agreed on next steps."
                }
            ],
            "action_items": [],
            "future_planning": {
                "open_issues": "None noted.",
                "next_meeting": "To be scheduled."
            }
        },
        "processed_sentences": processed,
        "main_concept_summary": "Meeting discussion recorded.",
        "off_topic_report": "No major off-topic digressions.",
        "action_items": []
    }


def generate_pdf_report(meeting_details: dict) -> bytes:
    # Generates native PDF report
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()

    indigo = colors.HexColor("#3730A3")
    dark_slate = colors.HexColor("#1E293B")
    soft_slate = colors.HexColor("#64748B")
    border_color = colors.HexColor("#CBD5E1")

    title_style = ParagraphStyle('DocTitle', parent=styles['Heading1'], fontSize=18, leading=22, textColor=indigo, spaceAfter=4, fontName='Helvetica-Bold')
    sub_style = ParagraphStyle('DocSub', parent=styles['Normal'], fontSize=9, leading=12, textColor=soft_slate, spaceAfter=12)
    h2_style = ParagraphStyle('SectionHeading', parent=styles['Heading2'], fontSize=12, leading=15, textColor=indigo, spaceBefore=10, spaceAfter=6, fontName='Helvetica-Bold')
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9, leading=13, textColor=dark_slate, spaceAfter=6)
    bold_body = ParagraphStyle('BoldBody', parent=body_style, fontName='Helvetica-Bold')

    story = []

    title = meeting_details.get("title", "Professional Meeting Report")
    m_id = meeting_details.get("meeting_id", "MTG")
    date_downloaded = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    story.append(Paragraph(f"🎙️ IntelliMin Type A Meeting Report: {title}", title_style))
    story.append(Paragraph(f"Meeting ID: <b>{m_id}</b> | Target Language: <b>{meeting_details.get('target_language', 'English')}</b> | Downloaded On: <b>{date_downloaded}</b>", sub_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=indigo, spaceAfter=10))

    type_a = meeting_details.get("type_a_report", {})
    meta = type_a.get("metadata", {})
    exec_sum = type_a.get("executive_summary", {})
    disc_pts = type_a.get("discussion_points", [])
    actions = type_a.get("action_items", meeting_details.get("action_items", []))
    future = type_a.get("future_planning", {})

    story.append(Paragraph("1. Document Metadata", h2_style))
    meta_table_data = [
        [Paragraph("<b>Meeting Title:</b>", bold_body), Paragraph(meta.get("meeting_title", title), body_style)],
        [Paragraph("<b>Date & Time:</b>", bold_body), Paragraph(f"{meta.get('date', meeting_details.get('date', ''))} at {meta.get('time', '')}", body_style)],
        [Paragraph("<b>Location / Link:</b>", bold_body), Paragraph(meta.get("location", "Virtual Conference"), body_style)],
        [Paragraph("<b>Facilitator / Note-Taker:</b>", bold_body), Paragraph(f"Facilitator: {meta.get('facilitator', 'Person 1 (Detected Voice)')} | Note-Taker: {meta.get('note_taker', 'IntelliMin')}", body_style)],
        [Paragraph("<b>Attendees:</b>", bold_body), Paragraph(", ".join(meta.get("attendees", ["Person 1 (Detected Voice)", "Person 2 (Detected Voice)"])), body_style)],
        [Paragraph("<b>Absentees:</b>", bold_body), Paragraph(", ".join(meta.get("absentees", ["None"])), body_style)],
    ]
    t_meta = Table(meta_table_data, colWidths=[140, 400])
    t_meta.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#F8FAFC')),
        ('GRID', (0,0), (-1,-1), 0.5, border_color),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('TOPPADDING', (0,0), (-1,-1), 3),
        ('BOTTOMPADDING', (0,0), (-1,-1), 3),
    ]))
    story.append(t_meta)
    story.append(Spacer(1, 8))

    story.append(Paragraph("2. Executive Summary", h2_style))
    story.append(Paragraph(f"<b>Objective:</b> {exec_sum.get('objective', meeting_details.get('main_summary', 'No objective recorded.'))}", body_style))
    story.append(Paragraph(f"<b>Outcome:</b> {exec_sum.get('outcome', 'Final agreement reached and tasks assigned.')}", body_style))
    story.append(Spacer(1, 8))

    story.append(Paragraph("3. Discussion Points (Agenda & Consensus)", h2_style))
    if disc_pts:
        disc_table_data = [[Paragraph("<b>Agenda Item</b>", bold_body), Paragraph("<b>Key Arguments</b>", bold_body), Paragraph("<b>Consensus / Agreement</b>", bold_body)]]
        for dp in disc_pts:
            disc_table_data.append([
                Paragraph(dp.get("agenda_item", "Agenda Topic"), body_style),
                Paragraph(dp.get("key_arguments", "-"), body_style),
                Paragraph(dp.get("consensus", "-"), body_style)
            ])
        t_disc = Table(disc_table_data, colWidths=[130, 230, 180])
        t_disc.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#EEF2FF')),
            ('GRID', (0,0), (-1,-1), 0.5, border_color),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('TOPPADDING', (0,0), (-1,-1), 3),
            ('BOTTOMPADDING', (0,0), (-1,-1), 3),
        ]))
        story.append(t_disc)
    else:
        story.append(Paragraph(meeting_details.get("main_summary", "No detailed discussion points recorded."), body_style))
    story.append(Spacer(1, 8))

    story.append(Paragraph("4. Action Items & Next Steps", h2_style))
    if actions:
        act_table_data = [[Paragraph("<b>Task Deliverable</b>", bold_body), Paragraph("<b>Owner</b>", bold_body), Paragraph("<b>Deadline</b>", bold_body), Paragraph("<b>Status</b>", bold_body)]]
        for a in actions:
            act_table_data.append([
                Paragraph(a.get("task", "-"), body_style),
                Paragraph(f"<code>{a.get('assignee', 'Person 1')}</code>", body_style),
                Paragraph(a.get("deadline", "Next Friday"), body_style),
                Paragraph(a.get("status", "Pending"), body_style)
            ])
        t_act = Table(act_table_data, colWidths=[230, 120, 100, 90])
        t_act.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#F0FDF4')),
            ('GRID', (0,0), (-1,-1), 0.5, border_color),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('TOPPADDING', (0,0), (-1,-1), 3),
            ('BOTTOMPADDING', (0,0), (-1,-1), 3),
        ]))
        story.append(t_act)
    else:
        story.append(Paragraph("<i>No action items recorded for this meeting.</i>", body_style))
    story.append(Spacer(1, 8))

    story.append(Paragraph("5. Future Planning & Next Session", h2_style))
    story.append(Paragraph(f"<b>Open / Tabled Issues:</b> {future.get('open_issues', 'None noted.')}", body_style))
    story.append(Paragraph(f"<b>Next Meeting Plan:</b> {future.get('next_meeting', 'To be announced.')}", body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_docx_report(meeting_details: dict) -> bytes:
    # Generates native Word DOCX
    doc = docx.Document()

    title = meeting_details.get("title", "Professional Meeting Report")
    m_id = meeting_details.get("meeting_id", "MTG")
    date_downloaded = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    doc.add_heading(f"🎙️ IntelliMin Type A Report: {title}", level=0)
    p_sub = doc.add_paragraph()
    p_sub.add_run(f"Meeting ID: {m_id} | Language: {meeting_details.get('target_language', 'English')} | Downloaded On: {date_downloaded}").italic = True

    type_a = meeting_details.get("type_a_report", {})
    meta = type_a.get("metadata", {})
    exec_sum = type_a.get("executive_summary", {})
    disc_pts = type_a.get("discussion_points", [])
    actions = type_a.get("action_items", meeting_details.get("action_items", []))
    future = type_a.get("future_planning", {})

    doc.add_heading("1. Document Metadata", level=1)
    t_meta = doc.add_table(rows=6, cols=2)
    t_meta.style = 'Table Grid'
    meta_rows = [
        ("Meeting Title:", meta.get("meeting_title", title)),
        ("Date & Time:", f"{meta.get('date', meeting_details.get('date', ''))} at {meta.get('time', '')}"),
        ("Location / Link:", meta.get("location", "Virtual Conference")),
        ("Facilitator / Note-Taker:", f"Facilitator: {meta.get('facilitator', 'Person 1 (Detected Voice)')} | Note-Taker: {meta.get('note_taker', 'IntelliMin')}"),
        ("Attendees:", ", ".join(meta.get("attendees", ["Person 1 (Detected Voice)", "Person 2 (Detected Voice)"]))),
        ("Absentees:", ", ".join(meta.get("absentees", ["None"])))
    ]
    for i, (label, val) in enumerate(meta_rows):
        t_meta.cell(i, 0).paragraphs[0].add_run(label).bold = True
        t_meta.cell(i, 1).paragraphs[0].text = str(val)

    doc.add_heading("2. Executive Summary", level=1)
    p_obj = doc.add_paragraph()
    p_obj.add_run("Objective: ").bold = True
    p_obj.add_run(exec_sum.get("objective", meeting_details.get("main_summary", "No objective recorded.")))

    p_out = doc.add_paragraph()
    p_out.add_run("Outcome: ").bold = True
    p_out.add_run(exec_sum.get("outcome", "Final results and agreement reached."))

    doc.add_heading("3. Discussion Points (Agenda & Consensus)", level=1)
    if disc_pts:
        t_disc = doc.add_table(rows=1, cols=3)
        t_disc.style = 'Table Grid'
        hdr_cells = t_disc.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run("Agenda Item").bold = True
        hdr_cells[1].paragraphs[0].add_run("Key Arguments").bold = True
        hdr_cells[2].paragraphs[0].add_run("Consensus / Agreement").bold = True
        for dp in disc_pts:
            row_cells = t_disc.add_row().cells
            row_cells[0].paragraphs[0].text = dp.get("agenda_item", "Agenda Topic")
            row_cells[1].paragraphs[0].text = dp.get("key_arguments", "-")
            row_cells[2].paragraphs[0].text = dp.get("consensus", "-")
    else:
        doc.add_paragraph(meeting_details.get("main_summary", "No detailed discussion points recorded."))

    doc.add_heading("4. Action Items & Next Steps", level=1)
    if actions:
        t_act = doc.add_table(rows=1, cols=4)
        t_act.style = 'Table Grid'
        hdr_act = t_act.rows[0].cells
        hdr_act[0].paragraphs[0].add_run("Task Deliverable").bold = True
        hdr_act[1].paragraphs[0].add_run("Owner").bold = True
        hdr_act[2].paragraphs[0].add_run("Deadline").bold = True
        hdr_act[3].paragraphs[0].add_run("Status").bold = True
        for a in actions:
            r_cells = t_act.add_row().cells
            r_cells[0].paragraphs[0].text = a.get("task", "-")
            r_cells[1].paragraphs[0].text = a.get("assignee", "Person 1")
            r_cells[2].paragraphs[0].text = a.get("deadline", "Next Friday")
            r_cells[3].paragraphs[0].text = a.get("status", "Pending")
    else:
        doc.add_paragraph("No action items recorded.")

    doc.add_heading("5. Future Planning & Next Session", level=1)
    p_fut1 = doc.add_paragraph()
    p_fut1.add_run("Open / Tabled Issues: ").bold = True
    p_fut1.add_run(future.get("open_issues", "None noted."))

    p_fut2 = doc.add_paragraph()
    p_fut2.add_run("Next Meeting Plan: ").bold = True
    p_fut2.add_run(future.get("next_meeting", "To be announced."))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_meeting_markdown_report(meeting_details: dict) -> str:
    # Generates Markdown report text
    title = meeting_details.get("title", "Meeting Report")
    m_id = meeting_details.get("meeting_id", "MTG")
    date_downloaded = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    type_a = meeting_details.get("type_a_report", {})
    meta = type_a.get("metadata", {})
    exec_sum = type_a.get("executive_summary", {})
    disc_pts = type_a.get("discussion_points", [])
    actions = type_a.get("action_items", meeting_details.get("action_items", []))
    future = type_a.get("future_planning", {})

    md = f"""# 🎙️ INTELLIMIN TYPE A MEETING REPORT: {title}
**Meeting ID:** `{m_id}`  
**Language:** {meeting_details.get('target_language', 'English')}  
**Report Download Date:** {date_downloaded}

---

## 1. Document Metadata
- **Meeting Title:** {meta.get('meeting_title', title)}
- **Date & Time:** {meta.get('date', meeting_details.get('date', ''))} at {meta.get('time', '')}
- **Location / Link:** {meta.get('location', 'Virtual Conference')}
- **Facilitator:** {meta.get('facilitator', 'Person 1 (Detected Voice)')}
- **Note-Taker:** {meta.get('note_taker', 'IntelliMin AI')}
- **Attendees:** {', '.join(meta.get('attendees', ['Person 1 (Detected Voice)', 'Person 2 (Detected Voice)']))}
- **Absentees:** {', '.join(meta.get('absentees', ['None']))}

---

## 2. Executive Summary
- **Objective:** {exec_sum.get('objective', meeting_details.get('main_summary', 'No objective recorded.'))}
- **Outcome:** {exec_sum.get('outcome', 'Final agreement reached.')}

---

## 3. Discussion Points (Agenda & Consensus)
"""
    if disc_pts:
        md += "| Agenda Item | Key Arguments | Consensus / Agreement |\n"
        md += "| :--- | :--- | :--- |\n"
        for dp in disc_pts:
            md += f"| {dp.get('agenda_item', 'Topic')} | {dp.get('key_arguments', '-')} | {dp.get('consensus', '-')} |\n"
    else:
        md += f"> {meeting_details.get('main_summary', 'No detailed discussion points recorded.')}\n"

    md += "\n---\n\n## 4. Action Items & Next Steps\n"
    if actions:
        md += "| Task Deliverable | Owner | Deadline | Status |\n"
        md += "| :--- | :--- | :--- |\n"
        for a in actions:
            md += f"| {a.get('task', '-')} | `{a.get('assignee', 'Person 1')}` | {a.get('deadline', 'Next Friday')} | {a.get('status', 'Pending')} |\n"
    else:
        md += "_No action items recorded for this meeting._\n"

    md += "\n---\n\n## 5. Future Planning & Next Session\n"
    md += f"- **Open / Tabled Issues:** {future.get('open_issues', 'None noted.')}\n"
    md += f"- **Next Meeting:** {future.get('next_meeting', 'To be announced.')}\n"

    md += "\n---\n*Type A Report automatically generated by IntelliMin.*"
    return md


# Initializes SQLite database tables
init_database()

# Loads separate style.css file
with open("style.css", "r", encoding="utf-8") as f:
    st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

# Loads separate header HTML
with open("header.html", "r", encoding="utf-8") as f:
    header_html = f.read()


# Initializes session state variables
if "session_counter" not in st.session_state:
    st.session_state["session_counter"] = 1

if "current_meeting_id" not in st.session_state:
    st.session_state["current_meeting_id"] = f"MTG-{datetime.now().strftime('%Y%m%d-%H%M%S')}"

if "live_utterances" not in st.session_state:
    st.session_state["live_utterances"] = []

if "editable_transcript" not in st.session_state:
    st.session_state["editable_transcript"] = ""

if "latest_generated_report" not in st.session_state:
    st.session_state["latest_generated_report"] = None

if "last_uploaded_filename" not in st.session_state:
    st.session_state["last_uploaded_filename"] = None

if "last_processed_mic_hash" not in st.session_state:
    st.session_state["last_processed_mic_hash"] = None


def reset_session_for_new_file():
    # Resets session completely clean
    st.session_state["current_meeting_id"] = f"MTG-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
    st.session_state["live_utterances"] = []
    st.session_state["editable_transcript"] = ""
    st.session_state["latest_generated_report"] = None
    st.session_state["last_uploaded_filename"] = None
    st.session_state["last_processed_mic_hash"] = None
    st.session_state["session_counter"] += 1


# Sidebar controls and settings
with st.sidebar:
    st.title("⚙️ Digital Controls")

    target_language = st.selectbox(
        "🌐 Target Output Language:",
        ["English", "Hindi", "Kannada", "Telugu", "Tamil", "Spanish", "French", "German", "Japanese", "Marathi", "Bengali"],
        index=0,
        help="Translates speech transcripts and full Type A report into requested language."
    )
    st.session_state["target_language"] = target_language

    st.markdown("---")
    if st.button("🔄 Start Fresh Meeting Session", use_container_width=True):
        reset_session_for_new_file()
        st.success("Session reset! All previous data cleared.")
        st.rerun()

    if st.button("🗑️ Delete Current Audio & Session Data", use_container_width=True):
        reset_session_for_new_file()
        st.success("Current audio and speech data deleted!")
        st.rerun()

# Renders header HTML banner
st.markdown(header_html, unsafe_allow_html=True)


# Top Summary Metrics Widget
db_meetings = search_meetings_db()
c_m1, c_m2, c_m3, c_m4 = st.columns(4)
with c_m1:
    st.markdown(f"""
        <div class="stat-card">
            <div class="label">Stored Database Meetings</div>
            <div class="number">{len(db_meetings)} / 10 Max</div>
        </div>
    """, unsafe_allow_html=True)

with c_m2:
    st.markdown(f"""
        <div class="stat-card">
            <div class="label">Target Output Language</div>
            <div class="number" style="font-size: 1.35rem; margin-top: 2px;">{target_language}</div>
        </div>
    """, unsafe_allow_html=True)

with c_m3:
    st.markdown(f"""
        <div class="stat-card">
            <div class="label">Report Standard</div>
            <div class="number" style="font-size: 1.15rem; margin-top: 4px; color: #0284C7 !important;">Type A Professional</div>
        </div>
    """, unsafe_allow_html=True)

with c_m4:
    st.markdown(f"""
        <div class="stat-card">
            <div class="label">PDF & DOCX Downloads</div>
            <div class="number" style="font-size: 1.15rem; margin-top: 4px; color: #10B981 !important;">Date Stamped</div>
        </div>
    """, unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)


# SEPARATE PANEL NAVIGATION TABS
panel1, panel2, panel3, panel4, panel5 = st.tabs([
    "🎙️ Panel 1: Voice & Media Input",
    "✏️ Panel 2: Transcript Editor & Evaluation",
    "📄 Panel 3: Type A Report & Downloads",
    "🗃️ Panel 4: Database Archives (10 Max)",
    "📊 Panel 5: Analytics & Flowchart"
])


# PANEL 1: WORKABLE VOICE RECORDING & MEDIA FILE UPLOAD
with panel1:
    st.markdown("""
        <div class="panel-card">
            <h3>🎙️ Panel 1: Workable Microphone Voice Recording & Media Upload</h3>
            <p>Every new recording session or file upload creates a fresh isolated session. Use Delete button to purge current audio data anytime.</p>
        </div>
    """, unsafe_allow_html=True)

    col_p1, col_p2 = st.columns(2)

    with col_p1:
        st.markdown("#### 🎤 Microphone Live Voice Input")
        
        c_mic1, c_mic2, c_mic3 = st.columns(3)
        with c_mic1:
            if st.button("🔴 Start New Recording", type="primary", use_container_width=True):
                reset_session_for_new_file()
                st.success("Started fresh recording session!")
                st.rerun()
        with c_mic2:
            if st.button("🗑️ Delete Audio Data", use_container_width=True):
                reset_session_for_new_file()
                st.success("Audio data deleted!")
                st.rerun()
        with c_mic3:
            if st.button("🧪 Load Sample", use_container_width=True):
                reset_session_for_new_file()
                sample_lines = [
                    "Person 1 (Detected Voice): Good morning team, let us review the cloud server infrastructure budget.",
                    "Person 2 (Detected Voice): We need to allocate $3000 for AWS server expansion by next Friday.",
                    "Person 1 (Detected Voice): By the way, did anyone try the new coffee machine on the 3rd floor? It makes great espresso!",
                    "Person 2 (Detected Voice): Yes, the coffee is awesome, but back to topic - who will handle the database migration?",
                    "David: I will take ownership of database migration and submit the plan by Wednesday."
                ]
                st.session_state["live_utterances"] = sample_lines
                st.session_state["editable_transcript"] = "\n".join(sample_lines)
                st.info("Loaded sample multi-speaker discussion!")

        current_counter = st.session_state["session_counter"]
        audio_val = st.audio_input("Record voice into microphone:", key=f"audio_mic_{current_counter}")
        if audio_val is not None:
            audio_val.seek(0)
            a_bytes = audio_val.read()
            import hashlib
            audio_hash = hashlib.md5(a_bytes).hexdigest() if a_bytes else None
            if audio_hash and st.session_state.get("last_processed_mic_hash") != audio_hash:
                st.session_state["last_processed_mic_hash"] = audio_hash
                with st.spinner("Transcribing spoken voice with Whisper..."):
                    tr_res = transcribe_audio_groq(a_bytes, filename="microphone_recording.wav")
                    spk_text = tr_res.get("text", "")
                    err_msg = tr_res.get("error", "")
                    if spk_text:
                        p_num = (len(st.session_state["live_utterances"]) % 3) + 1
                        spk_label = f"Person {p_num} (Detected Voice)"
                        line = f"{spk_label}: {spk_text}"
                        st.session_state["live_utterances"].append(line)
                        st.session_state["editable_transcript"] = "\n".join(st.session_state["live_utterances"])
                        st.success(f"Captured Speech: \"{spk_text}\" -> `{spk_label}`")
                    elif err_msg:
                        st.error(f"❌ Transcription error: {err_msg}")

        st.markdown("---")
        st.markdown("##### 🗣️ Add Spoken Voice Sentence")
        with st.form(f"add_voice_form_{current_counter}"):
            col_v1, col_v2 = st.columns([1, 3])
            with col_v1:
                spk_name_in = st.text_input("Speaker Label:", value="", help="Leave blank for Person 1 (Detected Voice)")
            with col_v2:
                spk_text_in = st.text_input("Spoken Voice Text:", value="", placeholder="Type or speak sentence...")
            btn_voice = st.form_submit_button("➕ Add Voice Sentence to Current Session")
            if btn_voice and spk_text_in.strip():
                if not spk_name_in.strip():
                    p_num = (len(st.session_state["live_utterances"]) % 3) + 1
                    spk_name_in = f"Person {p_num} (Detected Voice)"
                line = f"{spk_name_in.strip()}: {spk_text_in.strip()}"
                st.session_state["live_utterances"].append(line)
                st.session_state["editable_transcript"] = "\n".join(st.session_state["live_utterances"])
                st.success(f"Added voice sentence under `{spk_name_in}`!")

    with col_p2:
        st.markdown("#### 📁 Upload Audio or Video File")
        up_file = st.file_uploader(
            "Upload media file (.mp3, .wav, .m4a, .ogg, .flac, .mp4, .avi, .mov, .mkv):",
            type=["mp3", "wav", "m4a", "ogg", "flac", "mp4", "avi", "mov", "mkv"],
            key=f"file_uploader_{current_counter}"
        )

        if up_file is not None:
            if st.session_state["last_uploaded_filename"] != up_file.name:
                reset_session_for_new_file()
                st.session_state["last_uploaded_filename"] = up_file.name
                st.info(f"New file detected: `{up_file.name}`. Previous meeting session data cleared!")
                st.rerun()

            if st.button("⚡ Transcribe & Process Uploaded Media File", type="primary", use_container_width=True):
                with st.spinner(f"Transcribing '{up_file.name}' with Groq Whisper..."):
                    up_file.seek(0)
                    m_bytes = up_file.read()
                    tr_res = transcribe_audio_groq(m_bytes, filename=up_file.name)
                    spk_text = tr_res.get("text", "")
                    err_msg = tr_res.get("error", "")
                    if spk_text:
                        raw_sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', spk_text) if s.strip()]
                        if not raw_sentences:
                            raw_sentences = [spk_text.strip()]
                        formatted_lines = []
                        for idx, s_text in enumerate(raw_sentences):
                            p_num = (idx % 3) + 1
                            formatted_lines.append(f"Person {p_num} (Detected Voice): {s_text}")
                        st.session_state["live_utterances"] = formatted_lines
                        st.session_state["editable_transcript"] = "\n".join(formatted_lines)
                        st.success(f"Transcribed '{up_file.name}' ({len(formatted_lines)} sentences)! Switch to 'Panel 2' to review and evaluate report.")
                    elif err_msg:
                        st.error(f"❌ Transcription error: {err_msg}")

    if st.session_state["live_utterances"]:
        st.markdown("---")
        st.markdown("#### 📜 Transcribed Sentences in Current Session:")
        for idx, u in enumerate(st.session_state["live_utterances"], 1):
            st.markdown(f"**{idx}.** {u}")

        if st.button("🗑️ Clear / Delete Transcribed Sentences List"):
            reset_session_for_new_file()
            st.success("Transcribed sentences deleted!")
            st.rerun()


# PANEL 2: TRANSCRIPT EDITOR & REPORT EVALUATION
with panel2:
    st.markdown("""
        <div class="panel-card">
            <h3>✏️ Panel 2: Interactive Transcript Editor & Type A Evaluation</h3>
            <p>Review and edit transcribed sentences, assign speaker names or Person numbers, and trigger Type A report evaluation.</p>
        </div>
    """, unsafe_allow_html=True)

    m_title_input = st.text_input("Meeting Topic / Title:", value="Quarterly Infrastructure & Security Review")

    edited_text = st.text_area(
        "Spoken Utterances / Transcript (Edit line by line):",
        value=st.session_state["editable_transcript"],
        height=220,
        help="Edit text or speaker names before report evaluation."
    )
    st.session_state["editable_transcript"] = edited_text

    c_e1, c_e2 = st.columns(2)
    with c_e1:
        if st.button("🚀 Evaluate Transcript & Generate Type A Report", type="primary", use_container_width=True):
            raw_lines = [line.strip() for line in edited_text.split("\n") if line.strip()]
            if raw_lines:
                with st.spinner("Evaluating Type A Report, translating into " + target_language + ", and saving to SQLite database..."):
                    result = process_meeting_transcript_groq(raw_lines, target_language=target_language)
                    m_title = m_title_input.strip() or result.get("discussion_title", "Recorded Discussion")
                    m_id = st.session_state["current_meeting_id"]
                    date_str = datetime.now().strftime("%Y-%m-%d %H:%M")

                    save_meeting_to_db(
                        meeting_id=m_id,
                        title=m_title,
                        date_str=date_str,
                        target_lang=target_language,
                        utterances=result.get("processed_sentences", []),
                        summary=result.get("main_concept_summary", ""),
                        off_topic_rep=result.get("off_topic_report", ""),
                        action_items=result.get("action_items", []),
                        full_details=result
                    )

                    saved_details = get_meeting_details_db(m_id)
                    st.session_state["latest_generated_report"] = saved_details

                st.success("✅ Type A Professional Report evaluated and saved to database! Switch to 'Panel 3' to view and download PDF/DOCX.")
                st.balloons()
            else:
                st.warning("Please record speech or enter text in Panel 1 first.")

    with c_e2:
        if st.button("🗑️ Delete / Clear Current Session Data", use_container_width=True):
            reset_session_for_new_file()
            st.success("Cleared session data!")
            st.rerun()


# PANEL 3: TYPE A REPORT VIEWER & PDF/DOCX DOWNLOADS
with panel3:
    st.markdown("""
        <div class="panel-card">
            <h3>📄 Panel 3: Type A Professional Meeting Report & Exports</h3>
            <p>View the 5 structured Type A sections and download native PDF or Word documents with date stamps.</p>
        </div>
    """, unsafe_allow_html=True)

    if st.session_state.get("latest_generated_report"):
        rep = st.session_state["latest_generated_report"]
        type_a = rep.get("type_a_report", {})
        meta = type_a.get("metadata", {})
        exec_sum = type_a.get("executive_summary", {})
        disc_pts = type_a.get("discussion_points", [])
        actions = type_a.get("action_items", rep.get("action_items", []))
        future = type_a.get("future_planning", {})

        st.markdown(f"""
            <div class="typea-card">
                <h2>📄 Type A Professional Report: {rep.get('title', '')}</h2>
                <p>Meeting ID: <b>{rep.get('meeting_id')}</b> | Language: <b>{rep.get('target_language')}</b> | Date: <b>{rep.get('date')}</b></p>
            </div>
        """, unsafe_allow_html=True)

        st.markdown("#### 📥 Download Type A Report (Stamped with Date)")
        col_dl1, col_dl2, col_dl3 = st.columns(3)
        with col_dl1:
            pdf_bytes = generate_pdf_report(rep)
            st.download_button(
                label="📄 Download PDF Report (.pdf)",
                data=pdf_bytes,
                file_name=f"{rep.get('meeting_id')}_TypeA.pdf",
                mime="application/pdf",
                use_container_width=True
            )

        with col_dl2:
            docx_bytes = generate_docx_report(rep)
            st.download_button(
                label="📝 Download Word Report (.docx)",
                data=docx_bytes,
                file_name=f"{rep.get('meeting_id')}_TypeA.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

        with col_dl3:
            md_text = generate_meeting_markdown_report(rep)
            st.download_button(
                label="📋 Download Markdown (.md)",
                data=md_text,
                file_name=f"{rep.get('meeting_id')}_TypeA.md",
                mime="text/markdown",
                use_container_width=True
            )

        st.markdown("---")
        st.markdown("#### 1. Document Metadata")
        st.json(meta)

        st.markdown("#### 2. Executive Summary")
        st.markdown(f"**Objective:** {exec_sum.get('objective', rep.get('main_summary', ''))}")
        st.markdown(f"**Outcome:** {exec_sum.get('outcome', 'Agreement reached.')}")

        st.markdown("#### 3. Discussion Points (Agenda & Consensus)")
        if disc_pts:
            st.dataframe(pd.DataFrame(disc_pts), use_container_width=True, hide_index=True)
        else:
            st.markdown(f"> {rep.get('main_summary', '')}")

        st.markdown("#### 4. Action Items & Owners")
        if actions:
            st.dataframe(pd.DataFrame(actions), use_container_width=True, hide_index=True)

        st.markdown("#### 5. Future Planning")
        st.markdown(f"**Open / Tabled Issues:** {future.get('open_issues', 'None')}")
        st.markdown(f"**Next Meeting Schedule:** {future.get('next_meeting', 'TBD')}")
    else:
        st.info("No report generated in current session. Record voice in Panel 1 and evaluate in Panel 2!")


# PANEL 4: STORED DATABASE ARCHIVES (MAX 10 CAP)
with panel4:
    st.markdown("""
        <div class="panel-card">
            <h3>🗃️ Panel 4: Stored Database Archives (Strict 10 Cap Limit)</h3>
            <p>Search and view stored past meeting records in SQLite database. Older meetings are automatically rotated out to maintain 10 max cap.</p>
        </div>
    """, unsafe_allow_html=True)

    search_q = st.text_input("🔎 Search Stored Database (Topic, Title, Date):", value="")
    db_records = search_meetings_db(search_q)

    if db_records:
        st.markdown(f"Displaying **{len(db_records)}** stored meeting record(s) in SQLite database (Max 10 Cap).")
        for m in db_records:
            with st.expander(f"📌 {m['title']} | Date: {m['date']} | ID: {m['meeting_id']} | Language: {m['target_language']}", expanded=False):
                details = get_meeting_details_db(m["meeting_id"])
                
                col_d1, col_d2 = st.columns(2)
                with col_d1:
                    st.markdown("#### Executive Summary")
                    st.markdown(f"> {details.get('main_summary', '')}")
                with col_d2:
                    st.markdown("#### Off-Topic Digressions")
                    st.markdown(f"{details.get('off_topic_report', '')}")

                col_e1, col_e2, col_e3 = st.columns(3)
                with col_e1:
                    pdf_b = generate_pdf_report(details)
                    st.download_button(
                        label="📄 Download PDF",
                        data=pdf_b,
                        file_name=f"{m['meeting_id']}.pdf",
                        mime="application/pdf",
                        key=f"pdf_p4_{m['meeting_id']}"
                    )
                with col_e2:
                    docx_b = generate_docx_report(details)
                    st.download_button(
                        label="📝 Download DOCX",
                        data=docx_b,
                        file_name=f"{m['meeting_id']}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"docx_p4_{m['meeting_id']}"
                    )
                with col_e3:
                    if st.button("🗑️ Delete Record", key=f"del_p4_{m['meeting_id']}"):
                        if delete_meeting_db(m['meeting_id']):
                            st.success(f"Deleted {m['meeting_id']}")
                            st.rerun()
    else:
        st.info("No stored meetings in database.")


# PANEL 5: ACTION ITEMS, VISUAL ANALYTICS & PROJECT FLOWCHART
with panel5:
    st.markdown("""
        <div class="panel-card">
            <h3>📊 Panel 5: Master Action Items, Visual Analytics & Project Flowchart</h3>
            <p>Track tasks, view discussion category pie chart, speaker participation bar chart, and project architecture flowchart.</p>
        </div>
    """, unsafe_allow_html=True)

    all_stored = search_meetings_db()
    act_list = []
    all_transcripts = []
    for m in all_stored:
        d = get_meeting_details_db(m["meeting_id"])
        all_transcripts.extend(d.get("transcripts", []))
        for a in d.get("action_items", []):
            a["Meeting Title"] = d.get("title", "")
            a["Meeting ID"] = d.get("meeting_id", "")
            act_list.append(a)

    st.markdown("#### 📋 Action Items Tracker Table")
    if act_list:
        st.dataframe(pd.DataFrame(act_list), use_container_width=True, hide_index=True)
    else:
        st.info("No action items recorded.")

    st.markdown("---")
    st.markdown("#### 📊 Dual Visual Analytics Charts")
    col_g1, col_g2 = st.columns(2)

    with col_g1:
        st.markdown("##### 🥧 Discussion Category Breakdown (Pie Chart)")
        if all_transcripts:
            df_t = pd.DataFrame(all_transcripts)
            cat_counts = df_t["category"].value_counts().reset_index()
            cat_counts.columns = ["Category", "Count"]

            fig_pie = px.pie(
                cat_counts,
                names="Category",
                values="Count",
                color="Category",
                color_discrete_map={"Decision": "#10B981", "Action Item": "#2563EB", "Discussion": "#06B6D4", "Casual": "#F59E0B"},
                hole=0.35
            )
            fig_pie.update_layout(template="plotly_white")
            st.plotly_chart(fig_pie, use_container_width=True)
        else:
            st.info("No chart data available.")

    with col_g2:
        st.markdown("##### 📊 Speaker Participation Share (Bar Chart)")
        if all_transcripts:
            df_t = pd.DataFrame(all_transcripts)
            spk_counts = df_t["speaker"].value_counts().reset_index()
            spk_counts.columns = ["Speaker", "Count"]

            fig_bar = px.bar(
                spk_counts,
                x="Speaker",
                y="Count",
                color="Speaker",
                color_discrete_sequence=px.colors.qualitative.Pastel
            )
            fig_bar.update_layout(template="plotly_white", showlegend=False)
            st.plotly_chart(fig_bar, use_container_width=True)
        else:
            st.info("No chart data available.")
